# coding=utf-8
#读取docx中的文本代码示例
import docx
import random
from docx import Document
import os
import sys

import time
import wx



class PokerFrame(wx.Frame):
    def __init__(self, parent=None, id=-1, title="无线电英语通话自学软件"):
        wx.Frame.__init__(self, parent, id, title, size=(760,1010),
                          style=wx.DEFAULT_FRAME_STYLE ^ wx.RESIZE_BORDER)

        self.content_btn_id = 21
        self.style_btn_id = 23
        
        self.chapter    = 0   # 选择的序号
        self.new_knowledge = 1    # 学习个数
        self.y_n       = 'y'  # 是否继续
        self.output    = ""   # 输出的内容
        self.anser_str = ""   # 输入的答案

        self.max_n = 0
        self.knowledges = []   # 每次知识列表
        self.knowledge_dict = {} # 知识k,v 用于对比是否正确
        self.knowledge_key_list   = []  # 知识的key列表
        
        self.error_list    = []  # 本次回答错题

        ##### 检查部分
        self.check_dict = {}
        self.check_key_list = []
        self.score = 0 
        
        #### 示例
        self.doc_list = []
        
        ### ui 部分
        self.poker = wx.Panel(self, -1)
        self.panel = self.poker
        
        self.sizer = wx.BoxSizer(wx.VERTICAL)
        # wx.GridBagSizer(10, 20 )  # 列间隔为10，行间隔为20

        
        self.create_ui()

        self.show_chapter()

    def set_output(self, output):
        self.output_ctl.SetValue(output)

    def set_anser(self, anser_str):
        self.anser_ctl.SetValue(anser_str)
        pass 
        
    def show_chapter(self):
        t3_str = "\t"*5+"-"*50 + "\n"
        t3_str += "\t"*5+"欢迎来到无线电英语自学软件1.0\n"
        
        filename_short_list = get_filename()
        
        t3_str += "以下是可以学习的内容：\n"
        t3_str += "\n"
        t3_str += "\t 0 \t 思考题.docx \n"
        n = 0

        flist = sorted(filename_short_list)
        for i in range(len(flist)):
            n = i+1
            t3_str += "\t " + str(n) + "\t " + flist[i] + "\n"

        t3_str += "\t "+str(n+1)+"\t 第一章的章节测验\n"
        t3_str += "\t "+str(n+2)+"\t 第二章的章节测验\n"
        t3_str += "\t "+str(n+3)+"\t 第三章的章节测验\n"
        t3_str += "\t "+str(n+4)+"\t 第四章的章节测验\n"
        t3_str += "请根据序号选择你要学习的内容:\n"
        self.t3.SetValue(t3_str)
        self.max_n = n + 4
        pass
        
    ####  创建ui
    def create_ui(self):
        self.Center()    # 窗口剧中
        size = (110, 10)

        ########## 多行文本框 ##########
        hbox1 = wx.BoxSizer(wx.HORIZONTAL)        
        self.t3 = wx.TextCtrl(self.panel,size = (740,250),style = wx.TE_MULTILINE) 
        hbox1.Add(self.t3,1,wx.EXPAND|wx.ALIGN_LEFT|wx.ALL,5)
        #self.t3.Bind(wx.EVT_TEXT,self.OnKeyTyped)
        self.sizer.Add(hbox1)
        
        ### 选择数字
        hbox3 = wx.BoxSizer(wx.HORIZONTAL)
        number_lbl = wx.StaticText(self.panel,-1, label="请输入0到13:", size=size)
        hbox3.Add(number_lbl, 1, wx.EXPAND|wx.ALIGN_LEFT|wx.ALL, 15)

        self.number_ctl = wx.TextCtrl(self.panel, value=str(self.chapter), size=(150, 30))
        hbox3.Add(self.number_ctl, 2, wx.ALL, 10)
        self.sizer.Add(hbox3)
        

        ### 输入学习个数
        hbox4 = wx.BoxSizer(wx.HORIZONTAL)
        self.learn_num_lbl = wx.StaticText(self.panel,-1, label="学习个数:", size=size)
        hbox4.Add(self.learn_num_lbl, 1, wx.EXPAND|wx.ALIGN_LEFT|wx.ALL , 15)

        self.learn_num_ctl = wx.TextCtrl(self.panel, value=str(self.new_knowledge), size=(150, 30))
        hbox4.Add(self.learn_num_ctl, 2, wx.ALL, 10)

        
        tmp_lbl = wx.StaticText(self.panel, -1, label="", size=(50,10))
        hbox4.Add(tmp_lbl, 0.5, wx.EXPAND|wx.ALIGN_LEFT|wx.ALL, 10)
        self.start_btn = wx.Button(self.panel, self.style_btn_id, label=u"开始学习",  size=(120,30))

        self.start_btn.Bind(wx.EVT_BUTTON,self.onStart)
        
        hbox4.Add(self.start_btn, 1,  wx.EXPAND|wx.ALIGN_LEFT| wx.ALL, 10)
        
        self.sizer.Add(hbox4)

        ### 请输入Y/n
        hbox5 = wx.BoxSizer(wx.HORIZONTAL)
        learn_num_lbl = wx.StaticText(self.panel,-1, label="是否继续?(y/n):", size=size)
        hbox5.Add(learn_num_lbl, 1, wx.EXPAND|wx.ALIGN_LEFT|wx.ALL , 15)

        
        self.y_n_ctl = wx.TextCtrl(self.panel, value=self.y_n, size=(150, 30))
        hbox5.Add(self.y_n_ctl, 2, wx.ALL, 10)


        tmp_2_lbl = wx.StaticText(self.panel, -1, label="", size=(50,10))
        hbox5.Add(tmp_2_lbl, 0.5, wx.EXPAND|wx.ALIGN_LEFT|wx.ALL, 10)
        self.continue_btn = wx.Button(self.panel, self.style_btn_id, label=u"继续",  size=(120,30))

        self.continue_btn.Bind(wx.EVT_BUTTON,self.onContinue)
        
        hbox5.Add(self.continue_btn, 1,  wx.EXPAND|wx.ALIGN_LEFT| wx.ALL, 10)
        
        self.sizer.Add(hbox5)        


        ########## 输出
        hbox6 = wx.BoxSizer(wx.HORIZONTAL)        
        self.output_ctl = wx.TextCtrl(self.panel,size = (740,250),style = wx.TE_MULTILINE) 
        hbox6.Add(self.output_ctl, 1, wx.EXPAND|wx.ALIGN_LEFT|wx.ALL, 5)
        self.sizer.Add(hbox6)


        ######### 输入
        hbox7 = wx.BoxSizer(wx.HORIZONTAL)
        anser_lbl = wx.StaticText(self.panel,-1, label="请输入答案:", size=size)
        hbox7.Add(anser_lbl, 1, wx.EXPAND|wx.ALIGN_LEFT|wx.ALL , 15)
        
        self.anser_ctl = wx.TextCtrl(self.panel,size = (450,250),style = wx.TE_MULTILINE) 
        hbox7.Add(self.anser_ctl, 1, wx.EXPAND|wx.ALIGN_LEFT|wx.ALL, 5)


        self.commit_btn = wx.Button(self.panel, self.content_btn_id, label=u"提交",  size=(80,30))
        self.commit_btn.Bind(wx.EVT_BUTTON, self.check_anser)
        hbox7.Add(self.commit_btn, 1,  wx.EXPAND|wx.ALIGN_LEFT| wx.ALL, 10)        
        
        self.sizer.Add(hbox7)
        
        
        self.panel.SetSizerAndFit(self.sizer)


            
    ## 消息提示
    def showMsg(self, msg):
        dlg = wx.MessageDialog(None, u"%s"%(msg), u"提醒", wx.ICON_QUESTION)
        if dlg.ShowModal() == wx.ID_YES:
            self.Close(True)
        dlg.Destroy()
        return
        
    ## 退出
    def onQuit(self,evt):
        wx.Exit()


    ## 继续
    def onContinue(self, evt):
        self.y_n = self.y_n_ctl.GetValue()
        if self.y_n.strip() != "y" and self.y_n.strip() != "Y":
            self.showMsg("不继续学习,将退出软件!")
            self.onQuit(evt)
            pass
        else:
            i_chapter = self.number_ctl.GetValue()
            if int(i_chapter) == 0:
                self.display_example()
                self.check_example()
                
            if int(i_chapter) <= 9 and int(i_chapter) > 0 :
                self.display_know()
            if int(i_chapter) > 9 and int(i_chapter) <= self.max_n:
                self.check_write()
        pass 
        
    ## 启动处理
    def onStart(self, evt):
        self.knowledges         = []
        self.knowledge_key_list = []
        self.knowledge_dict     = {}
        self.error_list         = []
        self.check_key_list     = []
        self.check_dict         = {}
        self.score = 0 

        self.chapter = self.number_ctl.GetValue()

        self.y_n = self.y_n_ctl.GetValue()
        # print (self.chapter, self.new_knowledge, self.y_n)
        try:
            self.chapter = int(self.chapter)
           
        except Exception as e:
            self.showMsg(e)
            return 

        if self.y_n.strip() != "y" and self.y_n.strip() != "Y":
            self.onContinue(evt)
            return

        
        if int(self.chapter) == 0:
            self.display_chapter_0()
            pass 
        
        if int(self.chapter) <= 9 and int(self.chapter) > 0 :
            self.new_knowledge = self.learn_num_ctl.GetValue()
            try:
                self.new_knowledge = int(self.new_knowledge)
            except:
                self.showMsg("请输入学习个数是数字且大于0")
                return
            
            if int(self.chapter) > self.max_n:
                self.showMsg("请输入0到%d的数字"% self.max_n)
                return
            
            self.read_filename(int(self.chapter), int(self.new_knowledge))

        if int(self.chapter) > 9 and int(self.chapter) <= self.max_n:
            self.new_knowledge = 1
            self.test()
            self.display_test()
            pass 


    def display_chapter_0(self):
        self.doc_list = [{"filename":"起飞前与起飞阶段示例.docx",
                          "title":"请以起飞与起飞前阶段为话题，自行编写对"},
                         {"filename":"离场阶段示例.docx",
                          "title": "请以离场阶段为话题，自行编写对话"},
                         {"filename":"巡航阶段示例.docx",
                          "title": "请以巡航阶段为话题，自行编写对"},
                         {"filename":"进场阶段示例.docx",
                          "title": "请以进场阶段为话题，自行编写对话"},
                         {"filename":"最后进近及着陆阶段示例.docx",
                          "title": "请以最后进近及着陆为话题，自行编写对话"}]
        self.display_example()
        self.new_knowledge = len(self.doc_list)
        self.learn_num_ctl.SetValue(str(self.new_knowledge))
        pass
    
    def display_example(self):
        if len(self.doc_list) < 1:
            self.showMsg("写作结束!")
            return 
        ### doc1
        first_dict = self.doc_list[0]
        doc_name = first_dict.get("filename")
        whole_text1 = []
        try:
            doc11 = docx.Document(os.path.join('示例',doc_name))        
        except:
            self.showMsg("打开文件: %s 失败"%doc_name)
            return 
        paras1 = doc11.paragraphs
        for z in range(len(doc11.paragraphs)):
            whole_text1.append(doc11.paragraphs[z].text)
        outstr =  first_dict.get("title") + "\n"
        outstr += "以下为参考范例:\n " + "\n".join(whole_text1) + "\n"
        self.output_ctl.SetValue(outstr)

    ### 单词与对话
    def test(self):
        try:
            chapter = int(self.number_ctl.GetValue())
        except Exception as e:
            self.showMsg("请输入数字章节")
            return 
        document_1 = 0
        document_2 = 0
        if chapter == 10:
            document_1 = Document(os.path.join('单词和对话', '第一章单词.docx'))
            document_2 = Document(os.path.join('单词和对话', '第一章对话.docx'))
        if chapter == 11:
            document_1 = Document(os.path.join('单词和对话', '第二章单词.docx'))
            document_2 = Document(os.path.join('单词和对话', '第二章对话.docx'))
        if chapter == 12:
            document_1 = Document(os.path.join('单词和对话', '第三章单词.docx'))
            document_2 = Document(os.path.join('单词和对话', '第三章对话.docx'))
        if chapter == 13:
            document_1 = Document(os.path.join('单词和对话', '第四章单词.docx'))
            document_2 = Document(os.path.join('单词和对话', '第四章对话.docx'))
        if chapter != 13:
            document_1_list = []
            while True:
                a = random.randint(0, len(document_1.paragraphs))
                if a not in document_1_list and a % 2 == 0:
                    document_1_list.append(a)
                # 生产25个数的时候终止
                if len(document_1_list) == 25:
                    break
            document_2_list = []
            while True:
                b = random.randint(0, len(document_2.paragraphs))
                if b not in document_1_list and b % 2 == 0:
                    document_2_list.append(b)
                    # 生产25个数的时候终止
                if len(document_2_list) == 10:
                    break
            # print(document_1_list)
            # print(document_2_list)
            document_1_list_content = []
            document_2_list_content = []
            document_1_list_content_mean = []
            document_2_list_content_mean = []
            for i in range(len(document_1.paragraphs)):
                if i in document_1_list:
                    document_1_list_content.append(document_1.paragraphs[i].text)
                    document_1_list_content_mean.append(
                        document_1.paragraphs[i + 1].text)
            for j in range(len(document_2.paragraphs)):
                if j in document_1_list:
                    document_2_list_content.append(document_2.paragraphs[j].text)
                    document_2_list_content_mean.append(
                        document_2.paragraphs[j + 1].text)

        if chapter == 13:
            document_1_list = []
            while True:
                a = random.randint(0, len(document_1.paragraphs))
                if a not in document_1_list and a % 2 == 0:
                    document_1_list.append(a)
                # 生产25个数的时候终止
                if len(document_1_list) == 10:
                    break
            document_2_list = []
            while True:
                b = random.randint(0, len(document_2.paragraphs))
                if b not in document_1_list and b % 2 == 0:
                    document_2_list.append(b)
                # 生产25个数的时候终止
                if len(document_2_list) == 25:
                    break
            # print(document_1_list)
            # print(document_2_list)
            document_1_list_content = []
            document_2_list_content = []
            document_1_list_content_mean = []
            document_2_list_content_mean = []
            for i in range(len(document_1.paragraphs)):
                if i in document_1_list:
                    document_1_list_content.append(document_1.paragraphs[i].text)
                    document_1_list_content_mean.append(
                        document_1.paragraphs[i + 1].text)
            for j in range(len(document_2.paragraphs)):
                if j in document_1_list:
                    document_2_list_content.append(document_2.paragraphs[j].text)
                    document_2_list_content_mean.append(
                        document_2.paragraphs[j + 1].text)

        ######## 统一处理部分 ##########
        self.showMsg("---------------------------答题测验开始----------------------------")
            
        for i in range(len(document_1_list_content)):
            key = document_1_list_content[i]
            tmpstr = "根据英文单词默写出中文含义:\n" + str(key)
            self.check_key_list.append(key)
            self.check_dict[key] = {"mean": document_1_list_content_mean[i],
                                    "score": 4,
                                    "err_score": 0,
                                    "outstr": tmpstr}
        for i in range(len(document_2_list_content)):
            key = document_2_list_content[i]
            tmpstr = "根据英文句子写出中文含义:\n" + str(key)
            self.check_key_list.append(key)
            self.check_dict[key] = {"mean": document_2_list_content_mean[i],
                                    "score": 5,
                                    "err_score": 0,
                                    "outstr": tmpstr}
            #print (key)
        #### 打乱顺序
        random.shuffle(self.check_key_list)
        #### 设置学习个数
        # print (self.check_key_list)
        # print (len(self.check_key_list))
        self.new_knowledge = len(self.check_key_list)
        self.learn_num_ctl.SetValue(str(self.new_knowledge))
        return
        
        
    def read_filename(self, chapter, new_knowledge):
        try:
            a=os.listdir(os.path.join('单词和对话'))
        except:
            self.showMsg("没有找到'单词和对话'文件夹")
            return 
        filename_list = []
        filename_short_list = []
        for i in a:
            filename = os.path.join('单词和对话', i)
            filename_list.append(filename)
            filename_short_list.append(i)
        # print(filename_list)
        for l in range(len(filename_list)):
            if chapter-1 == l:
                self.get_content(filename_list[l],new_knowledge)        
        pass

    
    def get_content(self,j, new_knowledge):
        # 获取文档对象
        try:
            file=docx.Document(j)
        except Excaption as e:
            self.showMsg("打开文件失败: %s "%e)
            return 
        # print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
        #有多少行
        a=0
        #输出段落编号及段落内容
        for i in range(len(file.paragraphs)):
            a=a+1
        #存放英文
        list_1=[]
        #存放中文
        list_2=[]
        #row存放选择的行数
        row=[]
        while True:
            element=random.randint(0,a-1)
            if element%2==0 and element not in row:
                row.append(element)
                if len(row)==new_knowledge:
                    break
        for i in range(len(file.paragraphs)):
            if i in row:
                list_1.append(file.paragraphs[i].text)
                list_2.append(file.paragraphs[i+1].text)
        # print(list_1,list_2)
        #学习过程的处理
        for j in range(len(list_1)):
            tmpstr = "第{}个知识点:{}\n".format(j+1,list_1[j])
            tmpstr += "其中文含义为："+list_2[j] + "\n"
            self.knowledges.append(tmpstr)
            self.knowledge_dict[list_1[j]] = list_2[j]
        #### 将列表乱序
        self.knowledge_key_list = list(self.knowledge_dict.keys())
        random.shuffle(self.knowledge_key_list)
        self.display_know()
        return 
            
            
    ### check knowledges
    def check_knowledges(self):
        # print (self.knowledge_dict, self.knowledges)
        if len(self.knowledge_dict) < 1 or len(self.knowledge_key_list) < 1:
            if len(self.error_list) <= 0 :
                self.showMsg("检查结束,此次学习效果良好!")
            else:
                ret_str = "检查结束，共计错误{}次 \n".format(len(self.error_list))
                ret_str += "错误的知识有: \n"
                with open("错题本.txt", "w") as f:
                    for elist in self.error_list:
                        if len(elist) < 2:
                            continue
                        else:
                            f.write(elist[0])
                            f.write(elist[1])
                            ret_str += elist[0] + " : " + elist[1] + "\n"
                self.showMsg(ret_str)
            #### 结束
            return            
        else:
            key = self.knowledge_key_list[0]
            self.output_ctl.SetValue(u"请翻译 :  {}".format(key))
            pass
        
    ### 提交答案
    def check_anser(self, evt):
        try:
            i_chapter = int(self.number_ctl.GetValue())
        except Exception as e:
            self.showMsg("请输入章节为0到%d"%self.max_n)
            return
        if i_chapter == 0:
            self.check_example()
            
        if i_chapter > 0 and i_chapter <= 9:
            self.check_learn()

        if i_chapter > 9 :
            self.check_write()
            pass

    ### 检查0
    def check_example(self):
        if len(self.doc_list) < 1:
            self.showMsg("写作结束!")
            return
        input_str = self.anser_ctl.GetValue()
        if input_str != "":
            fname = "回答_" + self.doc_list[0].get("filename")
            with open(fname, "w") as f:
                f.write(input_str)
        ######  
        del(self.doc_list[0])
        self.display_example()
        self.new_knowledge -= 1
        self.learn_num_ctl.SetValue(str(self.new_knowledge))
        
        
    ## 检查章节检测
    def check_write(self):
        i_anser = self.anser_ctl.GetValue()
        if i_anser == "":
            self.showMsg("请输入答案")
            return
        ##############
        if len(self.check_dict) < 1 or len(self.check_key_list) < 1:
            error_n = len(self.error_list)
            self.showMsg("测试结束,你的得分为:%d,错误个数为:%d"%(self.score, error_n))
            return
        key = self.check_key_list[0]
        if i_anser == self.check_dict[key].get("mean"):
            cur_score = self.check_dict[key].get("score", 0)
            self.score += cur_score
        else:
            cur_score = self.check_dict[key].get("err_score", 0)
            self.score += cur_score
            self.error_list.append([key, self.check_dict[key].get("mean")])
        ####
        del(self.check_key_list[0])
        self.check_dict.pop(key)
        self.new_knowledge -= 1
        self.learn_num_ctl.SetValue(str(self.new_knowledge))
        self.display_test()

        pass 
        

    ## 检查单词与对话
    def check_learn(self):
        i_anser = self.anser_ctl.GetValue()
        if i_anser == "":
            self.showMsg("请输入回答!")
            return
        ##################
        if len(self.knowledge_key_list) < 1 or len(self.knowledge_dict) < 1:
            self.check_knowledges()
            return 
        
        key = self.knowledge_key_list[0]
        if i_anser != self.knowledge_dict[key]:
            self.showMsg("回答错误!")
            self.error_list.append([key, self.knowledge_dict[key]])
        else:
            self.showMsg("回答正确!")

        #### 无论正确与否都要清除第一个
        del(self.knowledge_key_list[0])
        self.knowledge_dict.pop(key)
        self.check_knowledges()

            
    ###### 显示
    def display_know(self):
        if len(self.knowledges) < 1:
            self.showMsg("新的知识点已经学习结束")
            if len(self.knowledge_dict) > 0:
                self.showMsg("下面开始检测")
                self.check_knowledges()
        else:
            self.output_ctl.SetValue(self.knowledges[0])
            del(self.knowledges[0])
            pass

    ####
    def display_test(self):
        if len(self.check_key_list) < 1:
            self.check_write()
            return
        key = self.check_key_list[0]
        outstr = self.check_dict[key].get("outstr", "")
        self.output_ctl.SetValue(outstr)
        pass 
    
    


def get_filename():
    try:
        a=os.listdir(os.path.join('单词和对话'))
    except:
        a = []
    filename_list=[]
    filename_short_list=[]
    for i in a:
        filename = os.path.join('单词和对话', i)
        filename_list.append(filename)
        filename_short_list.append(i)
    return filename_short_list



# 前端界面实现
def main():
    print ("main ...")
    app = wx.App(False)
    f = PokerFrame()
    f.Show()
    app.MainLoop()



if __name__ == '__main__':
    main()


